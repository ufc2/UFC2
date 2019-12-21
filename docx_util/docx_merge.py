import sys
sys.path.append("..")

from docx_operation import conflict_resolution as merge
from docx_operation import apply_sequence as apply

from make_hash import make_hash_paragraph

from docx import Document
import copy
import io

from utils import sha256

def delete_paragraph(paragraph, parent):
    p = paragraph._element
    parent.remove(p)
    paragraph._p = paragraph._element = None

def insert_paragraph(paragraph, parent, index):
    p = paragraph._element
    parent.insert(index, copy.deepcopy(p))

def save_docx(document):
    data = io.BytesIO()
    document.save(data)
    return data.getvalue()

def merge_docx(base, left, right):
    basedocx = Document(io.BytesIO(base))
    leftdocx = Document(io.BytesIO(left))
    rightdocx = Document(io.BytesIO(right))

    parent = basedocx.paragraphs[0]._element.getparent()

    paras = dict()
    codes2indexs = dict()
    indexs2codes = dict()
    code_index = 0

    def docx2str(document):
        dst_str = ''
        nonlocal code_index
        for para in document.paragraphs:
            hashcode = sha256(make_hash_paragraph(para).encode())
            if paras.get(hashcode) is None:
                paras[hashcode] = para
                codes2indexs[hashcode] = code_index
                dst_str += chr(code_index)
                code_index += 1
            else:
                dst_str += chr(codes2indexs[hashcode])
        return dst_str.encode()
    
    s_base = docx2str(basedocx)
    s_left = docx2str(leftdocx)
    s_right = docx2str(rightdocx)

    for code in codes2indexs:
        indexs2codes[codes2indexs[code]] = code

    sequence = merge(s_base, s_left, s_right)
    for op in sequence:
        if op.type == 'd':
            for index in op.data:
                para = paras[indexs2codes[index]]
                delete_paragraph(para, parent)
        else:
            for x in range(len(op.data) - 1, -1, -1):
                index = op.data[x]
                para = paras[indexs2codes[index]]
                insert_paragraph(para, parent, op.pos)

    return save_docx(basedocx)
